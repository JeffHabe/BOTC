"""
md_to_docx.py
將指定的 Markdown 文件轉換成 Word (.docx)，並嵌入其中的本機圖片。
用法：
    py scripts/md_to_docx.py guide/player_guide.md
輸出檔案將輸出在與輸入相同目錄下，副檔名改為 .docx。
"""
import sys
import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def fill_paragraph_with_inline(p, text: str, base_dir: Path):
    """將含行內 Markdown 與 HTML <img> 的 text 填充到現有段落 p 中，支援真正嵌入圖片。"""
    # 支援雙引號或單引號 src
    pattern = re.compile(r"(<img\s+[^>]*src=[\"']([^\"']+)[\"'][^>]*>)")
    parts = pattern.split(text)

    idx = 0
    while idx < len(parts):
        text_part = parts[idx]

        # 1. 處理文字部分 (含 markdown 樣式)
        if text_part:
            for seg_text, bold, code in parse_inline(text_part):
                run = p.add_run(seg_text)
                run.bold = bold
                if code:
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4F)

        # 2. 處理 <img> 圖片部分
        if idx + 1 < len(parts):
            img_src = parts[idx+2]

            # 將相對路徑轉為絕對路徑
            full_path = (base_dir / img_src).resolve()
            # 確保副檔名為 .png，防止殘留 svg 的問題
            img_path_str = str(full_path).replace(".svg", ".png")
            full_path_png = Path(img_path_str)

            if full_path_png.exists():
                try:
                    run = p.add_run()
                    # 插入行內圖片，高度限制在 Pt(12)，使其在 11pt 的微軟正黑體中能自然對齊
                    run.add_picture(str(full_path_png), height=Pt(12))
                except Exception as e:
                    p.add_run(f" [圖片載入失敗: {os.path.basename(img_src)}] ")
            else:
                p.add_run(f" [圖片未找到: {os.path.basename(img_src)}] ")

        idx += 3


def add_horizontal_rule(doc: Document):
    """在段落底部加一條水平分隔線（使用段落邊框）。"""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_run_bold(run, value: bool):
    run.bold = value


def parse_inline(text: str):
    """
    解析行內格式，回傳 list of (text, bold, code) tuples。
    支援 **bold**、`code`、![alt](url)（略過圖片，由行層處理）。
    """
    # 把圖片替換掉（行內圖片在外層單獨處理）
    text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
    segments = []
    # 解析 **bold** 和 `code`
    pattern = re.compile(r"(\*\*(.+?)\*\*|`([^`]+)`)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            segments.append((text[last : m.start()], False, False))
        if m.group().startswith("**"):
            segments.append((m.group(2), True, False))
        else:
            segments.append((m.group(3), False, True))
        last = m.end()
    if last < len(text):
        segments.append((text[last:], False, False))
    return segments


def add_paragraph_with_inline(doc: Document, text: str, base_dir: Path, style=None):
    """新增段落並套用行內格式，支援行內嵌入的 HTML <img> 圖片。"""
    p = doc.add_paragraph(style=style)
    fill_paragraph_with_inline(p, text, base_dir)
    return p


def add_image_from_md(doc: Document, alt: str, img_path: str, base_dir: Path):
    """嵌入圖片（若檔案存在）。"""
    # 解析相對路徑
    full_path = (base_dir / img_path).resolve()
    if full_path.exists():
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(full_path), width=Inches(2.75))
            # 圖說
            if alt:
                cap = doc.add_paragraph(alt)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].font.size = Pt(9)
                cap.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        except Exception as e:
            doc.add_paragraph(f"[圖片載入失敗: {img_path} - {e}]")
    else:
        doc.add_paragraph(f"[圖片未找到: {img_path}]")


def convert(md_file: str):
    md_path = Path(md_file).resolve()
    if md_path.suffix.lower() != ".md":
        print("錯誤：輸入檔案必須是 Markdown 檔案 (.md)")
        sys.exit(1)
    base_dir = md_path.parent
    out_path = md_path.with_suffix(".docx")

    doc = Document()

    # 設定預設中文字型
    style = doc.styles["Normal"]
    style.font.name = "微軟正黑體"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微軟正黑體")
    style.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    i = 0
    in_list = False  # 是否在無序清單中

    while i < len(lines):
        line = lines[i]

        # 空行
        if line.strip() == "":
            in_list = False
            i += 1
            continue

        # 水平線 ---
        if re.match(r"^-{3,}\s*$", line):
            add_horizontal_rule(doc)
            i += 1
            continue

        # 標題
        heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            title_text = heading_match.group(2).strip()
            heading_level = min(level, 4)  # docx 最多支援 Heading 4 以上
            p = doc.add_heading("", level=heading_level)
            fill_paragraph_with_inline(p, title_text, base_dir)
            in_list = False
            i += 1
            continue

        # 獨立圖片行 ![alt](path)
        img_match = re.match(r"^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$", line)
        if img_match:
            alt, img_path = img_match.group(1), img_match.group(2)
            add_image_from_md(doc, alt, img_path, base_dir)
            i += 1
            continue

        # 有序清單 1. 2. ...
        ordered_match = re.match(r"^\s*(\d+)\.\s+(.*)", line)
        if ordered_match:
            text = ordered_match.group(2).strip()
            add_paragraph_with_inline(doc, text, base_dir, style="List Number")
            in_list = True
            i += 1
            continue

        # 無序清單 * - 縮排
        unordered_match = re.match(r"^(\s*)[*\-]\s+(.*)", line)
        if unordered_match:
            indent = len(unordered_match.group(1))
            text = unordered_match.group(2).strip()
            list_style = "List Bullet 2" if indent >= 4 else "List Bullet"
            add_paragraph_with_inline(doc, text, base_dir, style=list_style)
            in_list = True
            i += 1
            continue

        # 一般段落（含行內圖片分離）
        in_list = False
        add_paragraph_with_inline(doc, line, base_dir)
        # 如果段落中有 Markdown 大圖片，依然在大圖行插入
        img_inline = re.findall(r"!\[([^\]]*)\]\(([^)]+)\)", line)
        for alt, img_path in img_inline:
            add_image_from_md(doc, alt, img_path, base_dir)

        i += 1

    doc.save(str(out_path))
    print(f"[完成] 已輸出：{out_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法：py scripts/md_to_docx.py <markdown檔案路徑>")
        sys.exit(1)
    convert(sys.argv[1])
